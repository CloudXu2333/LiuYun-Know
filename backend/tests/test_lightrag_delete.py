"""
测试 LightRAG 文档插入和删除功能

使用文档: 02、计算机学院关于开展2024-2025学年综合测评的通知.docx
"""
import os
import sys
import asyncio
from pathlib import Path

# 添加项目根目录到 Python 路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from lightrag import LightRAG, QueryParam
from lightrag.llm.openai import openai_complete_if_cache, openai_embed
from lightrag.utils import EmbeddingFunc
from lightrag.kg.shared_storage import initialize_pipeline_status

# 加载环境变量
from dotenv import load_dotenv
load_dotenv()

# 配置
WORKING_DIR = "./test_rag_storage"
TEST_DOC_PATH = "../02、计算机学院关于开展2024-2025学年综合测评的通知.docx"
FILE_NAME = "02、计算机学院关于开展2024-2025学年综合测评的通知.docx"

# API 配置
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_API_BASE = os.getenv("DEEPSEEK_API_BASE", "https://api.deepseek.com")
QWEN_API_KEY = os.getenv("QWEN_API_KEY")
QWEN_API_BASE = os.getenv("QWEN_API_BASE", "https://dashscope.aliyuncs.com/compatible-mode/v1")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-v3")


def extract_text_from_docx(file_path: str) -> str:
    """从 .docx 文件提取文本"""
    doc = Document(file_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragraphs.append(' | '.join(row_text))
    
    return '\n\n'.join(paragraphs)


async def create_llm_func(prompt, system_prompt=None, history_messages=[], **kwargs):
    """LLM 函数"""
    return await openai_complete_if_cache(
        model="deepseek-chat",
        prompt=prompt,
        system_prompt=system_prompt,
        history_messages=history_messages,
        api_key=DEEPSEEK_API_KEY,
        base_url=DEEPSEEK_API_BASE,
        **kwargs
    )


async def create_embedding_func(texts: list[str]):
    """Embedding 函数"""
    return await openai_embed(
        texts,
        model=EMBEDDING_MODEL,
        api_key=QWEN_API_KEY,
        base_url=QWEN_API_BASE,
    )


async def create_rag():
    """创建 LightRAG 实例"""
    os.makedirs(WORKING_DIR, exist_ok=True)
    
    embedding_func = EmbeddingFunc(
        embedding_dim=1024,
        func=create_embedding_func
    )
    
    rag = LightRAG(
        working_dir=WORKING_DIR,
        workspace="test_kb",
        graph_storage="Neo4JStorage",
        llm_model_func=create_llm_func,
        embedding_func=embedding_func,
        addon_params={
            "language": "Chinese",
        },
    )
    
    await rag.initialize_storages()
    await initialize_pipeline_status()
    
    return rag


async def test_insert():
    """测试插入文档"""
    print("=" * 60)
    print("📝 测试插入文档")
    print("=" * 60)
    
    # 提取文本
    doc_path = os.path.join(os.path.dirname(__file__), TEST_DOC_PATH)
    if not os.path.exists(doc_path):
        # 尝试从项目根目录查找
        doc_path = os.path.join(os.path.dirname(__file__), "..", "..", FILE_NAME)
    
    print(f"📄 文档路径: {doc_path}")
    print(f"   文件存在: {os.path.exists(doc_path)}")
    
    if not os.path.exists(doc_path):
        print("❌ 文档不存在，请检查路径")
        return None
    
    text = extract_text_from_docx(doc_path)
    print(f"📄 提取文本长度: {len(text)}")
    print(f"📄 文本预览: {text[:200]}...")
    
    # 创建 RAG 实例
    rag = await create_rag()
    
    # 插入文档，使用 file_paths 参数
    print(f"\n📥 插入文档，file_paths=['{FILE_NAME}']")
    await rag.ainsert(text, file_paths=[FILE_NAME])
    
    print("✅ 文档插入完成")
    
    return rag


async def test_query(rag):
    """测试查询"""
    print("\n" + "=" * 60)
    print("🔍 测试查询")
    print("=" * 60)
    
    query = "综合测评的时间安排是什么？"
    print(f"❓ 查询: {query}")
    
    result = await rag.aquery(
        query,
        param=QueryParam(mode="mix", only_need_context=True)
    )
    
    print(f"\n📋 查询结果:")
    print(result[:500] if len(result) > 500 else result)
    
    return result


async def test_list_docs(rag):
    """列出所有文档"""
    print("\n" + "=" * 60)
    print("📋 列出所有文档")
    print("=" * 60)
    
    # 读取 full_docs
    import json
    full_docs_path = os.path.join(WORKING_DIR, "test_kb", "kv_store_full_docs.json")
    if os.path.exists(full_docs_path):
        with open(full_docs_path, 'r', encoding='utf-8') as f:
            full_docs = json.load(f)
        print(f"📄 文档数量: {len(full_docs)}")
        for doc_id in full_docs.keys():
            print(f"   - {doc_id}")
    else:
        print("⚠️ full_docs 文件不存在")
    
    # 读取 text_chunks
    chunks_path = os.path.join(WORKING_DIR, "test_kb", "kv_store_text_chunks.json")
    if os.path.exists(chunks_path):
        with open(chunks_path, 'r', encoding='utf-8') as f:
            chunks = json.load(f)
        print(f"\n📄 分片数量: {len(chunks)}")
        for i, (chunk_id, chunk_info) in enumerate(list(chunks.items())[:3]):
            print(f"   [{i}] {chunk_id[:50]}...")
            print(f"       file_path: {chunk_info.get('file_path', 'N/A')}")
            print(f"       full_doc_id: {chunk_info.get('full_doc_id', 'N/A')[:50]}...")
    else:
        print("⚠️ text_chunks 文件不存在")
    
    return full_docs if os.path.exists(full_docs_path) else {}


async def test_delete(rag, doc_id: str = None):
    """测试删除文档"""
    print("\n" + "=" * 60)
    print("🗑️ 测试删除文档")
    print("=" * 60)
    
    if not doc_id:
        # 从 full_docs 获取第一个 doc_id
        import json
        full_docs_path = os.path.join(WORKING_DIR, "test_kb", "kv_store_full_docs.json")
        if os.path.exists(full_docs_path):
            with open(full_docs_path, 'r', encoding='utf-8') as f:
                full_docs = json.load(f)
            if full_docs:
                doc_id = list(full_docs.keys())[0]
    
    if not doc_id:
        print("⚠️ 没有找到可删除的文档")
        return
    
    print(f"🗑️ 删除文档: {doc_id}")
    
    try:
        await rag.adelete_by_doc_id(doc_id)
        print("✅ 文档删除完成")
    except Exception as e:
        print(f"❌ 删除失败: {e}")
        import traceback
        traceback.print_exc()


async def test_query_after_delete(rag):
    """删除后再次查询"""
    print("\n" + "=" * 60)
    print("🔍 删除后再次查询")
    print("=" * 60)
    
    query = "综合测评的时间安排是什么？"
    print(f"❓ 查询: {query}")
    
    try:
        result = await rag.aquery(
            query,
            param=QueryParam(mode="mix", only_need_context=True)
        )
        
        print(f"\n📋 查询结果:")
        if result and len(result.strip()) > 0:
            print(result[:500] if len(result) > 500 else result)
        else:
            print("✅ 没有查询到结果（符合预期，文档已删除）")
    except Exception as e:
        print(f"⚠️ 查询出错: {e}")


async def main():
    """主测试流程"""
    print("\n" + "=" * 60)
    print("🚀 LightRAG 文档插入和删除测试")
    print("=" * 60)
    
    # 1. 插入文档
    rag = await test_insert()
    if not rag:
        return
    
    # 2. 列出文档
    full_docs = await test_list_docs(rag)
    
    # 3. 查询测试
    await test_query(rag)
    
    # 4. 等待用户确认
    input("\n按 Enter 继续删除测试...")
    
    # 5. 删除文档
    if full_docs:
        doc_id = list(full_docs.keys())[0]
        await test_delete(rag, doc_id)
    
    # 6. 再次列出文档
    await test_list_docs(rag)
    
    # 7. 删除后查询
    await test_query_after_delete(rag)
    
    # 8. 关闭
    await rag.finalize_storages()
    
    print("\n" + "=" * 60)
    print("✅ 测试完成")
    print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
