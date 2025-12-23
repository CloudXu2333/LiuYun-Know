"""
Celery 文件处理任务
"""
import os
import tempfile
import asyncio
import numpy as np
from datetime import datetime
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker, Session
from lightrag import LightRAG, QueryParam
from lightrag.llm.openai import openai_complete_if_cache, openai_embed
from lightrag.utils import EmbeddingFunc
from lightrag.kg.shared_storage import initialize_pipeline_status
from app.core.celery_app import celery_app
from app.models.knowledge_base import KnowledgeFile, FileStatus
from app.services.minio_service import minio_service
from app.services.ocr_service import ocr_service
from app.config import settings


# 创建同步数据库引擎（用于 Celery 任务）
sync_database_url = settings.database_url.replace("sqlite+aiosqlite://", "sqlite://")
sync_engine = create_engine(sync_database_url, connect_args={"check_same_thread": False})
SyncSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=sync_engine)


def _get_workspace_label(working_dir: str) -> str:
    """从 working_dir 生成 Neo4j 的 workspace 标签"""
    import hashlib
    # 直接使用 working_dir 字符串生成哈希，不依赖当前工作目录
    # 统一使用正斜杠，确保跨平台一致性
    normalized_path = working_dir.replace('\\', '/')
    hash_value = hashlib.md5(normalized_path.encode()).hexdigest()[:16]
    print(f"[Workspace] working_dir={working_dir}, normalized={normalized_path}, workspace=kb_{hash_value}")
    return f"kb_{hash_value}"


async def _create_llm_func(prompt, system_prompt=None, history_messages=[], **kwargs):
    """创建 LLM 函数（使用 DeepSeek 模型进行实体抽取）"""
    return await openai_complete_if_cache(
        model="deepseek-chat",
        prompt=prompt,
        system_prompt=system_prompt,
        history_messages=history_messages,
        api_key=settings.deepseek_api_key,
        base_url=settings.deepseek_api_base,
        **kwargs
    )


async def _create_embedding_func(texts: list[str]) -> np.ndarray:
    """创建 Embedding 函数"""
    return await openai_embed(
        texts,
        model=settings.embedding_model,
        api_key=settings.qwen_api_key,
        base_url=settings.qwen_api_base,
    )


async def _insert_text_to_rag(working_dir: str, text: str, file_name: str = None):
    """
    在 Celery worker 中创建独立的 LightRAG 实例并插入文本
    注意：每次调用都创建新的 LightRAG 实例，避免 event loop 冲突
    
    Args:
        working_dir: 工作目录
        text: 要插入的文本
        file_name: 原始文件名（用于追踪来源）
    """
    os.makedirs(working_dir, exist_ok=True)
    
    # 检测现有的 workspace 目录，兼容旧版本数据
    workspace = None
    if os.path.exists(working_dir):
        for item in os.listdir(working_dir):
            item_path = os.path.join(working_dir, item)
            if os.path.isdir(item_path) and item.startswith('kb_'):
                workspace = item
                print(f"[Celery] 检测到现有 workspace: {workspace}")
                break
    
    # 如果没有找到现有目录，使用计算的 workspace 标签
    if not workspace:
        workspace = _get_workspace_label(working_dir)
        print(f"[Celery] 使用新 workspace: {workspace}")
    
    embedding_func = EmbeddingFunc(
        embedding_dim=1024,
        func=_create_embedding_func
    )
    
    # 每次创建新实例，避免 asyncio.Lock 绑定到旧的 event loop
    rag = LightRAG(
        working_dir=working_dir,
        workspace=workspace,
        graph_storage="Neo4JStorage",
        llm_model_func=_create_llm_func,
        embedding_func=embedding_func,
        addon_params={
            "language": "Chinese",
        },
    )
    
    try:
        await rag.initialize_storages()
        await initialize_pipeline_status()
        
        # 使用文件名作为 file_paths 参数，让 LightRAG 记录来源
        if file_name:
            await rag.ainsert(text, file_paths=[file_name])
            print(f"✅ 文本已插入到知识库: {working_dir}, workspace: {workspace}, 文件: {file_name}")
        else:
            await rag.ainsert(text)
            print(f"✅ 文本已插入到知识库: {working_dir}, workspace: {workspace}")
    finally:
        # 确保关闭存储连接，释放资源
        try:
            if hasattr(rag, 'chunk_entity_relation_graph') and rag.chunk_entity_relation_graph:
                await rag.chunk_entity_relation_graph.close()
        except Exception as e:
            print(f"⚠️ 关闭图存储时出错: {e}")
        
        # 清理 LightRAG 的共享锁，避免 event loop 冲突
        try:
            from lightrag.kg.shared_storage import _storage_locks
            # 清理该 workspace 相关的锁
            keys_to_remove = [k for k in _storage_locks.keys() if workspace in k]
            for key in keys_to_remove:
                del _storage_locks[key]
            if keys_to_remove:
                print(f"🧹 已清理 {len(keys_to_remove)} 个 storage locks")
        except Exception as e:
            print(f"⚠️ 清理 storage locks 时出错: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """
    从 .docx 文件提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        提取的文本内容
    """
    from docx import Document
    
    doc = Document(file_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragraphs.append(' | '.join(row_text))
    
    return '\n\n'.join(paragraphs)


def extract_text_from_doc(file_path: str) -> str:
    """
    从 .doc 文件提取文本（旧版 Word 格式）
    使用 antiword 或转换为 docx 后处理
    
    Args:
        file_path: 文件路径
        
    Returns:
        提取的文本内容
    """
    import subprocess
    import platform
    
    # 方法1：Windows 上使用 pywin32 (COM)
    if platform.system() == 'Windows':
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(file_path))
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                if text.strip():
                    return text
            finally:
                pythoncom.CoUninitialize()
        except Exception as e:
            print(f"⚠️ pywin32 处理 .doc 失败: {e}")
    
    # 方法2：使用 olefile 提取文本（纯 Python）
    try:
        import olefile
        
        ole = olefile.OleFileIO(file_path)
        # Word 文档的文本通常在 'WordDocument' 流中
        if ole.exists('WordDocument'):
            # 尝试读取文本
            stream = ole.openstream('WordDocument')
            data = stream.read()
            
            # 尝试解码文本（简单提取可见字符）
            text_parts = []
            try:
                # 尝试从二进制数据中提取文本
                decoded = data.decode('utf-16-le', errors='ignore')
                # 过滤掉控制字符，保留可打印字符
                text = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in decoded)
                text = ' '.join(text.split())  # 清理多余空格
                if len(text) > 100:  # 确保提取到了有意义的文本
                    ole.close()
                    return text
            except:
                pass
        ole.close()
    except Exception as e:
        print(f"⚠️ olefile 处理 .doc 失败: {e}")
    
    # 方法3：尝试使用 antiword（需要系统安装）
    try:
        result = subprocess.run(
            ['antiword', file_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    # 方法4：如果 antiword 不可用，尝试使用 LibreOffice 转换
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'docx', '--outdir', temp_dir, file_path],
                capture_output=True,
                timeout=120
            )
            if result.returncode == 0:
                # 找到转换后的文件
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                if os.path.exists(docx_path):
                    return extract_text_from_docx(docx_path)
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    raise ValueError("无法处理 .doc 文件。请安装 Microsoft Word 或将文件转换为 .docx 格式")


def extract_text_from_csv(file_path: str) -> str:
    """
    从 CSV 文件提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        提取的文本内容
    """
    import csv
    
    rows = []
    
    # 尝试不同的编码
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig', 'latin-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, newline='') as f:
                # 尝试检测分隔符
                sample = f.read(4096)
                f.seek(0)
                
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')
                except csv.Error:
                    dialect = csv.excel
                
                reader = csv.reader(f, dialect)
                
                for row in reader:
                    if any(cell.strip() for cell in row):
                        rows.append(' | '.join(cell.strip() for cell in row if cell.strip()))
                
                break  # 成功读取，退出循环
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    if not rows:
        raise ValueError("无法读取 CSV 文件，请检查文件编码")
    
    return '\n'.join(rows)


def extract_text_from_excel(file_path: str) -> str:
    """
    从 Excel 文件（.xlsx/.xls）提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        提取的文本内容
    """
    import openpyxl
    
    wb = openpyxl.load_workbook(file_path, data_only=True)
    all_text = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_rows = []
        
        # 添加工作表名称作为标题
        all_text.append(f"## 工作表: {sheet_name}")
        
        for row in sheet.iter_rows():
            row_values = []
            for cell in row:
                if cell.value is not None:
                    # 转换为字符串并清理
                    value = str(cell.value).strip()
                    if value:
                        row_values.append(value)
            
            if row_values:
                sheet_rows.append(' | '.join(row_values))
        
        if sheet_rows:
            all_text.extend(sheet_rows)
        
        all_text.append('')  # 工作表之间添加空行
    
    wb.close()
    return '\n'.join(all_text)


@celery_app.task(bind=True, name="process_file")
def process_file_task(self, file_id: int):
    """
    处理上传的文件
    
    Args:
        file_id: 文件 ID
    """
    # 使用同步方式处理
    return _process_file_sync(file_id)


def _process_file_sync(file_id: int):
    """同步处理文件（用于 Celery）"""
    db: Session = SyncSessionLocal()
    
    try:
        # 第一步：获取文件信息并更新状态为处理中
        file_record = db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).first()
        if not file_record:
            raise ValueError(f"文件记录不存在: {file_id}")
        
        # 更新状态为处理中
        file_record.status = FileStatus.PROCESSING
        db.commit()
        
        # 保存需要的信息
        original_filename = file_record.original_filename
        minio_path = file_record.minio_path
        file_type = file_record.file_type
        working_dir = file_record.knowledge_base.working_dir
        
        print(f"📝 开始处理文件: {original_filename}")
        
        # 第二步：处理文件
        with tempfile.TemporaryDirectory() as temp_dir:
            # 从 MinIO 下载文件
            temp_file_path = os.path.join(temp_dir, original_filename)
            minio_service.download_file(minio_path, temp_file_path)
            print(f"✅ 文件已下载: {temp_file_path}")
            
            # 提取文本内容
            extracted_text = ""
            
            if file_type.lower() == "pdf":
                # 使用 OCR 处理 PDF
                if ocr_service:
                    output_dir = os.path.join(temp_dir, "ocr_output")
                    extracted_text = ocr_service.process_pdf(temp_file_path, output_dir)
                else:
                    raise RuntimeError("OCR 服务未配置")
            
            elif file_type.lower() in ["txt", "md", "markdown"]:
                # 直接读取文本文件
                with open(temp_file_path, "r", encoding="utf-8") as f:
                    extracted_text = f.read()
            
            elif file_type.lower() == "docx":
                # 处理 .docx 文件
                extracted_text = extract_text_from_docx(temp_file_path)
            
            elif file_type.lower() == "doc":
                # 处理 .doc 文件（旧版 Word）
                extracted_text = extract_text_from_doc(temp_file_path)
            
            elif file_type.lower() == "csv":
                # 处理 CSV 文件
                extracted_text = extract_text_from_csv(temp_file_path)
            
            elif file_type.lower() in ["xlsx", "xls"]:
                # 处理 Excel 文件
                extracted_text = extract_text_from_excel(temp_file_path)
            
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
            
            if not extracted_text:
                raise ValueError("未能提取到文本内容")
            
            print(f"✅ 文本提取完成，长度: {len(extracted_text)}")
        
        # 第三步：插入到 LightRAG 知识库（使用新的事件循环）
        # 重要：清理全局状态，避免 asyncio.Lock 绑定到旧的 event loop
        try:
            # 清理可能存在的旧锁
            import lightrag.kg.shared_storage as shared_storage
            if hasattr(shared_storage, '_storage_locks'):
                shared_storage._storage_locks.clear()
        except Exception as e:
            print(f"⚠️ 清理存储锁时出错: {e}")
        
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            # 传递原始文件名，用于追踪来源
            loop.run_until_complete(_insert_text_to_rag(working_dir, extracted_text, original_filename))
            print(f"✅ 文本已插入知识库")
        finally:
            # 清理事件循环
            try:
                loop.run_until_complete(loop.shutdown_asyncgens())
            except Exception:
                pass
            loop.close()
            asyncio.set_event_loop(None)
        
        # 第四步：更新状态为完成
        file_record = db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).first()
        if file_record:
            file_record.status = FileStatus.COMPLETED
            file_record.processed_at = datetime.utcnow()
            file_record.error_message = None
            db.commit()
            
            print(f"✅ 文件处理完成: {original_filename}")
        
        return {"status": "success", "file_id": file_id}
    
    except Exception as e:
        print(f"❌ 文件处理失败: {str(e)}")
        
        # 更新状态为失败
        try:
            file_record = db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).first()
            if file_record:
                file_record.status = FileStatus.FAILED
                file_record.error_message = str(e)
                db.commit()
        except Exception as db_error:
            print(f"❌ 更新失败状态时出错: {str(db_error)}")
        
        raise
    
    finally:
        db.close()
